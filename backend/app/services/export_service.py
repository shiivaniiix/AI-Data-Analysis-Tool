import io
from datetime import datetime
from typing import Any

import matplotlib
matplotlib.use("Agg", force=True)
import matplotlib.pyplot as plt
from fastapi import HTTPException, status
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document

from app.services.ai_service import generate_ai_insights
from app.services.duckdb_service import (
    get_chart_data,
    get_suggested_charts_and_summary,
    get_table_preview,
)


def _safe_filename(name: str) -> str:
    # Keep it simple: remove path separators and weird chars.
    cleaned = "".join(ch for ch in name if ch.isalnum() or ch in ("-", "_", "."))
    return cleaned[:120] or "export"


def _to_png_bytes(fig) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


def render_chart_png(chart_kind: str, data: list[dict[str, Any]]) -> bytes:
    fig = plt.figure(figsize=(8, 4.5))

    if chart_kind == "bar":
        labels = [d["label"] for d in data][:20]
        values = [Number(d["value"]) for d in data][:20]
        plt.bar(labels, values)
        plt.xticks(rotation=35, ha="right")
        plt.tight_layout()
        return _to_png_bytes(fig)

    if chart_kind == "line":
        labels = [d["label"] for d in data][:30]
        values = [Number(d["value"]) for d in data][:30]
        plt.plot(range(len(values)), values, linewidth=2)
        plt.xticks(range(len(values)), labels, rotation=35, ha="right", fontsize=8)
        plt.tight_layout()
        return _to_png_bytes(fig)

    if chart_kind == "pie":
        labels = [d["name"] for d in data][:10]
        values = [Number(d["value"]) for d in data][:10]
        plt.pie(values, labels=labels, autopct="%1.1f%%")
        plt.tight_layout()
        return _to_png_bytes(fig)

    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported chart kind.")


def Number(x: Any) -> float:
    try:
        return float(x)
    except Exception:
        return 0.0


def _build_export_inputs(
    *,
    user_id: str,
    chat_id: str,
    filter_column: str | None,
    filter_values: list[str] | None,
) -> dict[str, Any]:
    suggested = get_suggested_charts_and_summary(user_id=user_id, chat_id=chat_id)
    summary = suggested["summary"]
    slicer = suggested["slicer"]
    inferred_filter_col = slicer.get("filter_column") if slicer else None
    resolved_filter_col = filter_column if filter_column is not None else inferred_filter_col
    resolved_filter_values = filter_values if filter_values is not None else []

    filter_active = bool(resolved_filter_col and resolved_filter_values)

    # Recompute each chart's data with active filter.
    updated_charts: list[dict[str, Any]] = []
    for c in suggested["suggested_charts"]:
        limit = 6 if c["chart_kind"] == "pie" else 10
        chart_data = get_chart_data(
            user_id=user_id,
            chat_id=chat_id,
            chart_kind=c["chart_kind"],
            x_column=c["x_column"],
            y_column=c["y_column"],
            filter_column=resolved_filter_col if filter_active else None,
            filter_values=resolved_filter_values if filter_active else None,
            limit=limit,
        )
        updated_charts.append({**c, data: chart_data})

    insights = generate_ai_insights(summary=summary, suggested_charts=updated_charts)

    preview = get_table_preview(
        user_id=user_id,
        chat_id=chat_id,
        filter_column=resolved_filter_col if filter_active else None,
        filter_values=resolved_filter_values if filter_active else None,
        preview_rows=25,
        max_columns=8,
    )

    return {
        "summary": summary,
        "charts": updated_charts,
        "insights": insights,
        "preview": preview,
        "filter": {"filter_column": resolved_filter_col, "filter_values": resolved_filter_values},
    }


def export_pdf_report(
    *,
    user_id: str,
    chat_id: str,
    filter_column: str | None = None,
    filter_values: list[str] | None = None,
) -> bytes:
    inputs = _build_export_inputs(
        user_id=user_id,
        chat_id=chat_id,
        filter_column=filter_column,
        filter_values=filter_values,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
    styles = getSampleStyleSheet()
    title_style = styles["Title"]
    normal = styles["BodyText"]

    elements: list[Any] = []
    elements.append(Paragraph("DataChat AI Report", title_style))
    elements.append(Spacer(1, 10))
    elements.append(Paragraph(f"Chat ID: <b>{chat_id}</b>", normal))
    elements.append(Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", normal))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("AI Summary", styles["Heading2"]))
    for ins in inputs["insights"]:
        elements.append(Paragraph(f"• {ins}", normal))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Summary Statistics", styles["Heading2"]))
    summary_rows: list[list[Any]] = []
    summary_rows.append(["Rows", inputs["summary"].get("row_count", "")])
    summary_rows.append(["Columns", ", ".join(inputs["summary"].get("columns", [])[:10])])
    numeric_stats = inputs["summary"].get("numeric_stats", {})
    if isinstance(numeric_stats, dict):
        for k, v in list(numeric_stats.items())[:5]:
            mean = v.get("mean", "") if isinstance(v, dict) else ""
            summary_rows.append([f"Mean({k})", mean])

    tbl = Table(summary_rows, hAlign="LEFT", colWidths=[140, 320])
    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ]
        )
    )
    elements.append(tbl)
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Charts", styles["Heading2"]))
    for c in inputs["charts"]:
        img_bytes = render_chart_png(c["chart_kind"], c["data"])
        img = RLImage(io.BytesIO(img_bytes), width=460, height=260)
        elements.append(Paragraph(c["title"], styles["Heading3"]))
        elements.append(img)
        elements.append(Spacer(1, 12))

    elements.append(Paragraph("Data Preview", styles["Heading2"]))
    preview = inputs["preview"]
    cols = preview["columns"][:8]
    rows = preview["rows"][:25]
    table_data: list[list[Any]] = [cols]
    for r in rows:
        row_out: list[Any] = []
        for j in range(len(cols)):
            row_out.append(r[j] if j < len(r) else "")
        table_data.append(row_out)
    data_tbl = Table(table_data, hAlign="LEFT", repeatRows=1)
    data_tbl.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
            ]
        )
    )
    elements.append(data_tbl)

    doc.build(elements)
    buf.seek(0)
    return buf.read()


def export_docx_report(
    *,
    user_id: str,
    chat_id: str,
    filter_column: str | None = None,
    filter_values: list[str] | None = None,
) -> bytes:
    inputs = _build_export_inputs(
        user_id=user_id,
        chat_id=chat_id,
        filter_column=filter_column,
        filter_values=filter_values,
    )

    doc = Document()
    doc.add_heading("DataChat AI Report", level=0)
    doc.add_paragraph(f"Chat ID: {chat_id}")
    doc.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    doc.add_heading("AI Summary", level=1)
    for ins in inputs["insights"]:
        doc.add_paragraph(ins, style="List Bullet")

    doc.add_heading("Summary Statistics", level=1)
    doc.add_paragraph(f"Rows: {inputs['summary'].get('row_count', '')}")
    doc.add_paragraph(
        f"Columns: {', '.join(inputs['summary'].get('columns', [])[:10])}"
    )

    doc.add_heading("Charts", level=1)
    for c in inputs["charts"]:
        doc.add_heading(c["title"], level=2)
        img_bytes = render_chart_png(c["chart_kind"], c["data"])
        img_stream = io.BytesIO(img_bytes)
        # python-docx requires a filename-like object; BytesIO works.
        doc.add_picture(img_stream, width=6.0)

    doc.add_heading("Data Preview", level=1)
    preview = inputs["preview"]
    cols = preview["columns"][:8]
    rows = preview["rows"][:25]

    table = doc.add_table(rows=1, cols=len(cols))
    hdr_cells = table.rows[0].cells
    for idx, col in enumerate(cols):
        hdr_cells[idx].text = str(col)

    for r in rows:
        cells = table.add_row().cells
        for idx in range(len(cols)):
            cells[idx].text = str(r[idx]) if idx < len(r) else ""

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

